from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
from pathlib import Path

def set_margins(section, top=0.45, bottom=0.45, left=0.55, right=0.55):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_font(run, size=9, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_font(r, size=9, bold=True)

def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("• " + text)
    set_font(r, size=8.5)

def add_role(doc, heading, subheading, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(heading)
    set_font(r, size=9, bold=True)
    r = p.add_run(" | " + subheading)
    set_font(r, size=8.5)

    for b in bullets:
        add_bullet(doc, b)

def render_resume(resume, out_path):
    doc = Document()
    set_margins(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Douglas Gordon Daly")
    set_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Portland, OR | douglasgdaly@gmail.com | (650) 586-9720 | linkedin.com/in/douglasdaly | github.com/dougdaly")
    set_font(r, size=8)

    # Ensure header.title is present and validated
    header = resume.get("header", {})
    # Prefer explicit header.title, fall back to top-level target_title
    title = header.get("title") or resume.get("target_title")
    if not title:
        raise ValueError(
            "Missing required header.title (or target_title). Add a header.title field to target_resume_v1.json"
        )
    # If header didn't include title but we fell back, persist it for consistency
    header["title"] = title
    resume["header"] = header

    # Target title (from header.title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, size=10, bold=True)

    # Summary
    add_section_heading(doc, "Summary")
    for item in resume["summary"]:
        add_bullet(doc, item)

    # Expertise / technologies
    add_section_heading(doc, "Core Expertise")
    p = doc.add_paragraph()
    r = p.add_run(" | ".join(resume["core_expertise"]))
    set_font(r, size=8.5)

    add_section_heading(doc, "Technologies")
    p = doc.add_paragraph()
    r = p.add_run(" | ".join(resume["technologies"]))
    set_font(r, size=8.2)

    # Selected experience
    add_section_heading(doc, "Selected Experience")
    for role in resume["selected_experience"]:
        add_role(
            doc,
            role["heading"],
            role["subheading"],
            role.get("bullets", [])
        )

    # Additional experience
    add_section_heading(doc, "Additional Experience")
    for role in resume["additional_experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f'{role["heading"]} | {role["subheading"]}')
        set_font(r, size=8.5)

    # Selected projects
    add_section_heading(doc, "Selected Projects")
    for project in resume["selected_projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(project["heading"])
        set_font(r, size=8.5, bold=True)
        for b in project.get("bullets", []):
            add_bullet(doc, b)

    # Education
    add_section_heading(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run(" | ".join(resume["education"]))
    set_font(r, size=8.2)

    # Honors
    add_section_heading(doc, "Honors & Awards")
    for item in resume["honors_awards"]:
        add_bullet(doc, item)

    doc.save(out_path)

with open("artifacts/target_resume_v1.json") as f:
    target_resume_v1 = json.load(f)

# If the canonical artifact lacks header.title, prefer target_title from positioning.
header = target_resume_v1.get("header", {}) or {}
if not header.get("title") and target_resume_v1.get("target_title"):
    header["title"] = target_resume_v1.get("target_title")
    target_resume_v1["header"] = header

# Transform canonical resume schema into the simplified renderer shape expected
# by `render_resume()` so we don't have to change the canonical artifact format.
def find_section(resume, heading_name):
    for s in resume.get("sections", []):
        if (s.get("heading") or "").strip().lower() == heading_name.lower():
            return s
    # fallback: match by substring
    for s in resume.get("sections", []):
        if heading_name.lower() in (s.get("heading") or "").lower():
            return s
    return None

def extract_summary(resume):
    s = find_section(resume, "summary")
    if not s:
        return []
    if s.get("type") == "paragraph":
        return s.get("content", [])
    return [str(x) for x in s.get("content", [])]

def extract_core_expertise(resume):
    s = find_section(resume, "core expertise") or find_section(resume, "core_expertise")
    if not s:
        return []
    if isinstance(s.get("content"), list):
        return [str(x) for x in s.get("content")]
    return [str(s.get("content"))]

def extract_technologies(resume):
    s = find_section(resume, "technologies")
    if not s:
        return []
    return [str(x) for x in s.get("content", [])]

def extract_professional_experience(resume):
    s = find_section(resume, "professional experience")
    out = []
    if not s:
        return out
    for exp in s.get("content", []):
        heading = exp.get("role") or exp.get("heading") or ""
        organization = exp.get("organization", "")
        dates = exp.get("dates", "")
        subheading = " | ".join(x for x in [organization, dates] if x)

        bullets = []
        if exp.get("type") == "bullet":
            bullets = [str(x) for x in exp.get("content", [])]
        elif exp.get("type") == "subsections":
            for block in exp.get("content", []):
                if isinstance(block, dict):
                    for b in block.get("content", []):
                        bullets.append(str(b))
                elif isinstance(block, str):
                    bullets.append(block)
        else:
            # fallback: try to extract strings from content
            for item in exp.get("content", []):
                if isinstance(item, str):
                    bullets.append(item)
                elif isinstance(item, dict):
                    for b in item.get("content", []):
                        bullets.append(str(b))

        out.append({
            "heading": heading,
            "subheading": subheading,
            "bullets": bullets,
        })

    return out

def extract_selected_projects(resume):
    s = find_section(resume, "selected projects")
    out = []
    if not s:
        return out
    for p in s.get("content", []):
        if isinstance(p, dict):
            out.append({
                "heading": p.get("heading") or p.get("label") or "",
                "bullets": p.get("content", []) if isinstance(p.get("content", []), list) else [p.get("content", "")],
            })
        elif isinstance(p, str):
            out.append({"heading": p, "bullets": []})
    return out

resume_for_renderer = {
    "header": target_resume_v1.get("header", {}),
    "target_title": target_resume_v1.get("header", {}).get("title") or target_resume_v1.get("target_title"),
    "summary": extract_summary(target_resume_v1),
    "core_expertise": extract_core_expertise(target_resume_v1),
    "technologies": extract_technologies(target_resume_v1),
    "selected_experience": extract_professional_experience(target_resume_v1),
    "additional_experience": [],
    "selected_projects": extract_selected_projects(target_resume_v1),
    "education": [],
    "honors_awards": [],
}

render_resume(
    resume_for_renderer,
    "artifacts/douglas_daly_principal_ai_engineer_v1.docx",
)