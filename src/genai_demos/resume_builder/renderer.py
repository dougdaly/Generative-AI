"""Generic resume renderer.

Consumes a renderer-neutral resume JSON file and a YAML layout file.

Schema direction:
- Top-level resume has `header` and `sections`.
- Each section has `heading`, `type`, and `content`.
- Supported block types are intentionally small:
  paragraph, bullet, inline_list, subsections, experience.
- Experience items use role/organization/dates plus optional type/content.
- Subsection items use label/context/dates plus optional type/content.

"""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any
from .helpers import load_json, load_yaml
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}

# Helpers
def add_keep_together_cell(doc: Document):
    """Add a one-cell table row that LibreOffice should keep together."""
    table = doc.add_table(rows=1, cols=1)
    set_table_borders_none(table)
    set_table_indent(table, 0)
    set_row_cant_split(table.rows[0])

    cell = table.cell(0, 0)
    set_cell_margins(cell, 0)

    return cell

def set_table_borders_none(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)

    tblPr.append(borders)


def set_row_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def add_bullet_paragraph(
    container,
    text: str,
    layout: dict[str, Any],
    style_name: str | None = "bullet",
):
    '''The key to bullet alignment. Both normal bullets and table-wrapped 
    first bullets go through the exact same code path.'''
    cfg = style_cfg(layout, style_name)
    bullet_cfg = layout.get("bullets", {})

    p = container.add_paragraph(style="List Bullet")
    apply_paragraph_style(p, cfg)

    p.paragraph_format.left_indent = Inches(float(bullet_cfg.get("left_indent", 0.18)))
    p.paragraph_format.first_line_indent = Inches(
        -float(bullet_cfg.get("hanging_indent", 0.12))
    )

    run = p.add_run(str(text))
    apply_run_style(run, cfg)

    return p


def add_heading_with_first_bullet(
    doc: Document,
    heading_text: str,
    first_bullet: str,
    layout: dict[str, Any],
    heading_style_name: str | None,
    bullet_style_name: str | None,
) -> None:
    '''Use a one-cell table only when we need heading + first bullet to stay together.'''
    cell = add_keep_together_cell(doc)

    heading_p = cell.paragraphs[0]
    heading_cfg = style_cfg(layout, heading_style_name)
    apply_paragraph_style(heading_p, heading_cfg)

    heading_run = heading_p.add_run(str(heading_text))
    apply_run_style(heading_run, heading_cfg)

    add_bullet_paragraph(cell, first_bullet, layout, bullet_style_name)

def set_table_indent(table, indent_twips: int = 0) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblInd = tblPr.find(qn("w:tblInd"))

    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)

    tblInd.set(qn("w:w"), str(indent_twips))
    tblInd.set(qn("w:type"), "dxa")


def set_cell_margins(cell, margin_twips: int = 0) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for margin_name in ("top", "start", "bottom", "end", "left", "right"):
        node = tcMar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tcMar.append(node)

        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def style_cfg(layout: dict[str, Any], style_name: str | None) -> dict[str, Any]:
    if not style_name:
        return {}
    return layout.get("styles", {}).get(style_name, {}) or {}


def type_cfg(layout: dict[str, Any], type_name: str | None) -> dict[str, Any]:
    if not type_name:
        return {}
    return layout.get("types", {}).get(type_name, {}) or {}


def set_document_defaults(doc: Document, layout: dict[str, Any]) -> None:
    margins = layout.get("page", {}).get("margins", {})
    section = doc.sections[0]
    section.top_margin = Inches(float(margins.get("top", 0.5)))
    section.bottom_margin = Inches(float(margins.get("bottom", 0.5)))
    section.left_margin = Inches(float(margins.get("left", 0.6)))
    section.right_margin = Inches(float(margins.get("right", 0.6)))

    fonts = layout.get("fonts", {})
    doc.styles["Normal"].font.name = fonts.get("base", "Aptos")
    doc.styles["Normal"].font.size = Pt(float(fonts.get("body_size", 9)))


def add_bottom_border(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def apply_paragraph_style(paragraph, cfg: dict[str, Any]) -> None:
    if "align" in cfg:
        paragraph.alignment = ALIGN_MAP.get(
            str(cfg["align"]).lower(),
            WD_ALIGN_PARAGRAPH.LEFT,
        )

    paragraph.paragraph_format.space_before = Pt(float(cfg.get("space_before", 0)))
    paragraph.paragraph_format.space_after = Pt(float(cfg.get("space_after", 0)))

    if "keep_with_next" in cfg:
        paragraph.paragraph_format.keep_with_next = bool(cfg["keep_with_next"])

    if "keep_together" in cfg:
        paragraph.paragraph_format.keep_together = bool(cfg["keep_together"])

    if "page_break_before" in cfg:
        paragraph.paragraph_format.page_break_before = bool(cfg["page_break_before"])

    if cfg.get("border_bottom"):
        add_bottom_border(paragraph)


def apply_run_style(run, cfg: dict[str, Any]) -> None:
    # Missing bools intentionally default to False.
    run.bold = bool(cfg.get("bold", False))
    run.italic = bool(cfg.get("italic", False))
    run.underline = bool(cfg.get("underline", False))
    if "size" in cfg:
        run.font.size = Pt(float(cfg["size"]))

def apply_keep_with_next(paragraph) -> None:
    if paragraph is not None:
        paragraph.paragraph_format.keep_with_next = True


def add_text_paragraph(doc: Document, text: str, cfg: dict[str, Any]):
    if not text:
        return None

    p = doc.add_paragraph()
    apply_paragraph_style(p, cfg)

    run = p.add_run(text)
    apply_run_style(run, cfg)

    return p


def add_labeled_inline_paragraph(
    doc: Document,
    label_text: str,
    items: list[Any],
    layout: dict[str, Any],
    label_style_name: str | None,
    inline_cfg: dict[str, Any],
):
    values = [str(item) for item in items or [] if item]

    if not label_text and not values:
        return None

    label_cfg = style_cfg(layout, label_style_name)
    item_cfg = style_cfg(layout, inline_cfg.get("content_style"))

    p = doc.add_paragraph()
    apply_paragraph_style(p, item_cfg or label_cfg)

    if label_text:
        r = p.add_run(label_text + ": ")
        apply_run_style(r, label_cfg)

    if values:
        r = p.add_run(str(inline_cfg.get("separator", " | ")).join(values))
        apply_run_style(r, item_cfg)

    return p

def apply_keep_together(paragraph) -> None:
    paragraph.paragraph_format.keep_together = True


def render_header(doc: Document, resume: dict[str, Any], layout: dict[str, Any]) -> None:
    header = resume.get("header", {}) or {}
    for row in layout.get("header", {}).get("rows", []):
        values: list[str] = []
        for field in row.get("fields", []):
            value = header.get(field)
            if value:
                values.append(str(value))
        if values:
            text = str(row.get("separator", "")).join(values)
            add_text_paragraph(doc, text, style_cfg(layout, row.get("style")))


def render_section_heading(doc: Document, section: dict[str, Any], layout: dict[str, Any]) -> None:
    heading = section.get("heading", "")
    if not heading:
        return
    cfg = dict(style_cfg(layout, "section_heading"))
    if cfg.get("uppercase", True):
        heading = heading.upper()
    add_text_paragraph(doc, str(heading), cfg)


def format_field_value(value: Any, fmt: str | None) -> str:
    text = str(value)
    if fmt:
        return fmt.replace("{value}", text)
    return text


def format_object(obj: dict[str, Any], formatter_name: str, layout: dict[str, Any]) -> tuple[str, str | None]:
    formatter = layout.get("formatters", {}).get(formatter_name, {}) or {}
    field_formats = formatter.get("field_formats", {}) or {}
    parts: list[str] = []
    for field in formatter.get("fields", []):
        value = obj.get(field)
        if value:
            parts.append(format_field_value(value, field_formats.get(field)))
    return str(formatter.get("separator", "")).join(parts), formatter.get("style")


def render_bullet(
    doc: Document,
    items: list[Any],
    layout: dict[str, Any],
    style_name: str | None = "bullet",
) -> None:
    for item in items or []:
        if item:
            add_bullet_paragraph(doc, str(item), layout, style_name)

def render_paragraph(doc: Document, items: list[Any] | str, layout: dict[str, Any], style_name: str | None) -> None:
    cfg = style_cfg(layout, style_name)
    if isinstance(items, str):
        items = [items]
    for item in items or []:
        if item:
            add_text_paragraph(doc, str(item), cfg)


def render_inline_list(doc: Document, items: list[Any], layout: dict[str, Any], cfg: dict[str, Any]) -> None:
    values = [str(item) for item in items or [] if item]
    if not values:
        return
    text = str(cfg.get("separator", " | ")).join(values)
    add_text_paragraph(doc, text, style_cfg(layout, cfg.get("content_style")))


def render_items(doc: Document, items: list[dict[str, Any]], layout: dict[str, Any], cfg: dict[str, Any]) -> None:
    formatter_name = cfg.get("item_formatter")
    default_child_type = cfg.get("default_child_type")
    pagination = layout.get("pagination", {})
    keep_with_next_styles = set(pagination.get("keep_with_next_styles", []))

    for item in items or []:
        child_type = item.get("type") or default_child_type
        child_content = item.get("content")

        text = ""
        style_name = None

        if formatter_name:
            text, style_name = format_object(item, formatter_name, layout)

        if text and child_type == "bullet" and isinstance(child_content, list) and child_content:
            bullet_cfg = type_cfg(layout, "bullet")
            bullet_style_name = bullet_cfg.get("content_style")

            add_heading_with_first_bullet(
                doc=doc,
                heading_text=text,
                first_bullet=str(child_content[0]),
                layout=layout,
                heading_style_name=style_name,
                bullet_style_name=bullet_style_name,
            )

            remaining_bullets = child_content[1:]
            if remaining_bullets:
                render_bullet(doc, remaining_bullets, layout, bullet_style_name)

            continue

        # Common compact pattern: subsection label plus an inline list on the same line.
        # This is self-contained, so continue is correct here.
        if child_type == "inline_list" and child_content is not None:
            inline_cfg = type_cfg(layout, "inline_list")
            paragraph = add_labeled_inline_paragraph(
                doc,
                text,
                child_content,
                layout,
                style_name,
                inline_cfg,
            )

            if style_name in keep_with_next_styles:
                apply_keep_with_next(paragraph)

            continue

        heading_paragraph = None

        if text:
            heading_paragraph = add_text_paragraph(
                doc,
                text,
                style_cfg(layout, style_name),
            )

            if style_name in keep_with_next_styles:
                apply_keep_with_next(heading_paragraph)

        # Important: do not remove this.
        # This renders the bullets/subsections under each experience heading.
        if child_type and child_content is not None:
            render_block(doc, child_type, child_content, layout)

def render_block(doc: Document, block_type: str, content: Any, layout: dict[str, Any]) -> None:
    cfg = type_cfg(layout, block_type)
    if not cfg:
        raise ValueError(f"Unsupported type: {block_type!r}")

    if block_type == "paragraph":
        render_paragraph(doc, content, layout, cfg.get("content_style"))
    elif block_type == "bullet":
        render_bullet(doc, content, layout, cfg.get("content_style"))
    elif block_type == "inline_list":
        render_inline_list(doc, content, layout, cfg)
    elif block_type in {"experience", "subsections"}:
        render_items(doc, content, layout, cfg)
    else:
        raise ValueError(f"Unsupported type: {block_type!r}")

def build_document(resume: dict[str, Any], layout: dict[str, Any]) -> Document:
    doc = Document()
    set_document_defaults(doc, layout)
    render_header(doc, resume, layout)

    for section in resume.get("sections", []):
        render_section_heading(doc, section, layout)
        render_block(doc, section.get("type"), section.get("content", []), layout)

    return doc


def render_resume_docx(resume: dict[str, Any], layout: dict[str, Any], output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_document(resume, layout)
    doc.save(output_path)

    return output_path


def render_resume_pdf(
    resume: dict[str, Any],
    layout: dict[str, Any],
    output_path: str | Path,
    *,
    soffice_path: str | None = None,
    keep_intermediate_docx: bool = False,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if output_path.suffix.lower() != ".pdf":
        raise ValueError(f"PDF output path must end in .pdf: {output_path}")

    if keep_intermediate_docx:
        docx_path = output_path.with_suffix(".docx")
        render_resume_docx(resume, layout, docx_path)
        generated_pdf = convert_docx_to_pdf(docx_path, soffice_path=soffice_path)

        if generated_pdf.resolve() != output_path.resolve():
            shutil.move(str(generated_pdf), str(output_path))

        return output_path

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx_path = Path(tmp_dir) / f"{output_path.stem}.docx"
        render_resume_docx(resume, layout, tmp_docx_path)

        generated_pdf = convert_docx_to_pdf(tmp_docx_path, soffice_path=soffice_path)
        shutil.copy2(generated_pdf, output_path)

    return output_path


def render_resume(
    resume: dict[str, Any],
    layout: dict[str, Any],
    output_path: str | Path,
    *,
    soffice_path: str | None = None,
    keep_intermediate_docx: bool = False,
) -> Path:
    output_path = Path(output_path)
    suffix = output_path.suffix.lower()

    if suffix == ".docx":
        return render_resume_docx(resume, layout, output_path)

    if suffix == ".pdf":
        return render_resume_pdf(
            resume,
            layout,
            output_path,
            soffice_path=soffice_path,
            keep_intermediate_docx=keep_intermediate_docx,
        )

    raise ValueError(
        f"Unsupported output format: {output_path.suffix}. Use .docx or .pdf."
    )


def convert_docx_to_pdf(docx_path: str | Path, soffice_path: str | None = None) -> Path:
    docx_path = Path(docx_path).resolve()
    soffice = soffice_path or shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nolockcheck",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
    )
    return docx_path.with_suffix(".pdf")

